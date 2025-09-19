import streamlit as st
import pandas as pd
import requests
import json
import re
from io import BytesIO
from docx import Document
import os

st.set_page_config(page_title="GitLab Dashboard", layout="wide")

# ------------------------
# Helper Functions
# ------------------------
def safe_get_json(response):
    try:
        if response and response.status_code == 200:
            return response.json()
        return []
    except Exception:
        return []

def safe_milestone(issue):
    milestone = issue.get("milestone")
    if milestone and isinstance(milestone, dict):
        return milestone.get("title", "")
    return ""

def normalize_label(raw_label):
    return re.sub(r"^\d+\-?\s*", "", raw_label.strip())

def parse_labels(labels):
    parsed = {"Team": [], "Status": [], "Sprint": [], "Project": [], "Workstream": [], "Milestone": []}
    for raw_label in labels:
        label = normalize_label(raw_label)
        if "::" not in label:
            continue
        key, value = label.split("::",1)
        key = key.strip().title()
        value = value.strip()
        if key.startswith("Team"): parsed["Team"].append(value)
        elif key.startswith("Status"): parsed["Status"].append(value)
        elif key.startswith("Sprint"): parsed["Sprint"].append(value)
        elif key.startswith("Project"): parsed["Project"].append(value)
        elif key.startswith("Workstream"): parsed["Workstream"].append(value)
        elif key.startswith("Milestone"): parsed["Milestone"].append(value)
    return parsed

def fetch_issues(base_url, project_id, token):
    headers = {"PRIVATE-TOKEN": token}
    url = f"{base_url}/api/v4/projects/{project_id}/issues?per_page=100"
    resp = requests.get(url, headers=headers, verify=False)
    issues = safe_get_json(resp)
    parsed = []
    for issue in issues:
        labels = issue.get("labels", [])
        parsed_labels = parse_labels(labels)
        sprints = parsed_labels["Sprint"] or ["None"]
        teams = parsed_labels["Team"] or ["None"]
        for sprint in sprints:
            for team in teams:
                parsed.append({
                    "iid": issue.get("iid"),
                    "title": issue.get("title","Untitled"),
                    "description": issue.get("description",""),
                    "team": team,
                    "status": ", ".join(parsed_labels["Status"]) if parsed_labels["Status"] else "",
                    "sprint": sprint,
                    "project": ", ".join(parsed_labels["Project"]) if parsed_labels["Project"] else "",
                    "workstream": ", ".join(parsed_labels["Workstream"]) if parsed_labels["Workstream"] else "",
                    "milestone": ", ".join(parsed_labels["Milestone"]) if parsed_labels["Milestone"] else safe_milestone(issue),
                    "labels": labels,
                    "web_url": issue.get("web_url","")
                })
    return pd.DataFrame(parsed)

def update_issue(base_url, project_id, token, iid, payload):
    headers = {"PRIVATE-TOKEN": token}
    url = f"{base_url}/api/v4/projects/{project_id}/issues/{iid}"
    resp = requests.put(url, headers=headers, json=payload, verify=False)
    return resp.status_code == 200

STATUS_COLORS = {
    "Done": "#2ecc71",
    "In Progress": "#f1c40f",
    "To Do": "#e74c3c",
    "Blocked": "#e67e22",
    "": "#bdc3c7"
}

# ------------------------
# Sidebar for Connection & Filters
# ------------------------
st.sidebar.header("GitLab Connection")
base_url = st.sidebar.text_input("Base URL", value="https://gitlab.com")
project_id = st.sidebar.text_input("Project ID")
token = st.sidebar.text_input("Access Token", type="password")

if not (base_url and project_id and token):
    st.warning("Enter Base URL, Project ID, and Access Token to continue.")
    st.stop()

df = fetch_issues(base_url, project_id, token)
if df.empty:
    st.warning("No issues found. Check project ID or token.")
    st.stop()

# Filters
st.sidebar.header("Filters")
all_sprints = sorted(df["sprint"].replace("", "None").unique())
all_teams = sorted(df["team"].replace("", "None").unique())
all_projects = sorted(df["project"].replace("", "None").unique())
all_milestones = sorted(df["milestone"].replace("", "None").unique())
all_statuses = sorted(df["status"].replace("", "None").unique())

filter_sprint = st.sidebar.multiselect("Sprint", all_sprints)
filter_team = st.sidebar.multiselect("Team", all_teams)
filter_project = st.sidebar.multiselect("Project", all_projects)
filter_milestone = st.sidebar.multiselect("Milestone", all_milestones)
filter_status = st.sidebar.multiselect("Status", all_statuses)

if st.sidebar.button("Reset Filters"):
    filter_sprint = filter_team = filter_project = filter_milestone = filter_status = []

filtered_df = df.copy()
if filter_sprint: filtered_df = filtered_df[filtered_df["sprint"].isin(filter_sprint)]
if filter_team: filtered_df = filtered_df[filtered_df["team"].isin(filter_team)]
if filter_project: filtered_df = filtered_df[filtered_df["project"].isin(filter_project)]
if filter_milestone: filtered_df = filtered_df[filtered_df["milestone"].isin(filter_milestone)]
if filter_status: filtered_df = filtered_df[filtered_df["status"].isin(filter_status)]

# ------------------------
# Tabs
# ------------------------
tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["Overview","Kanban","Hygiene","Commentary","Edit"]
)

# ------------------------
# Tab 1: Overview
# ------------------------
with tab1:
    st.subheader("Overview Dashboard")
    col1,col2,col3,col4 = st.columns(4)
    for idx, field in enumerate(["status","team","sprint","project"]):
        col = [col1,col2,col3,col4][idx]
        with col:
            st.markdown(f"#### {field.title()}")
            values = sorted(filtered_df[field].replace("", "None").unique())
            for val in values:
                count = len(filtered_df[filtered_df[field]==val])
                color = STATUS_COLORS.get(val,"#3498db")
                if st.button(f"{val} ({count})", key=f"overview_{field}_{val}"):
                    st.experimental_set_query_params(**{field:val})

    st.markdown("### Issues Table")
    st.dataframe(filtered_df[["team","title","description","status","project","web_url"]])
    st.download_button("⬇️ Download CSV", filtered_df.to_csv(index=False), "overview.csv")

# ------------------------
# Tab 2: Kanban
# ------------------------
with tab2:
    st.subheader("Kanban Board (Status → Columns, Team → Grouping)")
    statuses = sorted(filtered_df["status"].replace("", "None").unique())
    board_col, detail_col = st.columns([3,2])
    with board_col:
        st.markdown('<div style="display:flex; overflow-x:auto; gap:15px;">', unsafe_allow_html=True)
        for status in statuses:
            lane_issues = filtered_df[filtered_df["status"]==status].sort_values("team")
            st.markdown(f'<div style="flex:0 0 300px; background:#f9f9f9; padding:10px; border-radius:5px;">', unsafe_allow_html=True)
            st.markdown(f"### {status}")
            for _, row in lane_issues.iterrows():
                color = STATUS_COLORS.get(row["status"],"#3498db")
                if st.button(f"{row['title']} ({row['team']})", key=f"kanban_{row['iid']}"):
                    st.session_state["selected_issue"] = row["iid"]
            st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with detail_col:
        iid = st.session_state.get("selected_issue")
        if iid:
            row = filtered_df[filtered_df["iid"]==iid].iloc[0]
            st.markdown(f"### {row['title']}")
            st.markdown(f"- **Team:** {row['team']}")
            st.markdown(f"- **Status:** {row['status']}")
            st.markdown(f"- **Sprint:** {row['sprint']}")
            st.markdown(f"- **Project:** {row['project']}")
            st.markdown(f"- **Workstream:** {row['workstream']}")
            st.markdown(f"- **Milestone:** {row['milestone']}")
            st.markdown(f"- **Web URL:** [{row['web_url']}]({row['web_url']})")

# ------------------------
# Tab 3: Hygiene
# ------------------------
with tab3:
    st.subheader("Hygiene Checks")
    hygiene_fields = ["team","status","sprint","project","milestone","title"]
    hygiene_counts = {field: len(filtered_df[filtered_df[field]==""]) for field in hygiene_fields}
    cols = st.columns(len(hygiene_fields))
    for idx, field in enumerate(hygiene_fields):
        col = cols[idx]
        with col:
            if st.button(f"{field.title()}: {hygiene_counts[field]}", key=f"hygiene_{field}"):
                st.dataframe(filtered_df[filtered_df[field]==""])
                # Quick Fix
                fix_field = st.selectbox("Set value for field", [""]+list(filtered_df[field].unique()))
                if st.button("Apply Fix"):
                    for _, row in filtered_df[filtered_df[field]==""].iterrows():
                        iid = row["iid"]
                        payload_labels = row["labels"]
                        if field.title() in ["Team","Status","Sprint","Project","Milestone"]:
                            payload_labels.append(f"{field.title()}::{fix_field}")
                        payload = {"labels": payload_labels}
                        update_issue(base_url, project_id, token, iid, payload)
                    st.success(f"Applied fix to {field}")

# ------------------------
# Tab 4: Commentary
# ------------------------
commentary_file = "commentary.json"
if os.path.exists(commentary_file):
    with open(commentary_file,"r") as f:
        commentary_data = json.load(f)
else:
    commentary_data = {}

with tab4:
    st.subheader("Commentary")
    all_sprints = sorted(filtered_df["sprint"].unique())
    selected_sprint = st.selectbox("Select Sprint", all_sprints)
    commentary = commentary_data.get(selected_sprint, {})
    commentary["Scope"] = st.text_area("Scope", commentary.get("Scope",""))
    commentary["Key Dates"] = st.text_area("Key Dates", commentary.get("Key Dates",""))
    commentary["Achievements"] = st.text_area("Achievements", commentary.get("Achievements",""))
    commentary["Next Steps"] = st.text_area("Next Steps", commentary.get("Next Steps",""))
    commentary["Challenges"] = st.text_area("Challenges", commentary.get("Challenges",""))
    if st.button("Save Commentary"):
        commentary_data[selected_sprint] = commentary
        with open(commentary_file,"w") as f:
            json.dump(commentary_data,f,indent=2)
        st.success("Saved successfully!")

    # Download as Word
    if st.button("Download Commentary as Word"):
        doc = Document()
        doc.add_heading(f"Commentary for {selected_sprint}",0)
        for key, val in commentary.items():
            doc.add_heading(key, level=1)
            doc.add_paragraph(val)
        buffer = BytesIO()
        doc.save(buffer)
        st.download_button("⬇️ Download", buffer.getvalue(), f"commentary_{selected_sprint}.docx")

# ------------------------
# Tab 5: Edit
# ------------------------
with tab5:
    st.subheader("Edit Issues")
    issue_titles = [f"#{row['iid']} - {row['title']}" for _, row in filtered_df.iterrows()]
    selected_issue_title = st.selectbox("Select Issue", [""]+issue_titles)
    if selected_issue_title:
        iid = int(selected_issue_title.split(" ")[0][1:])
        row = filtered_df[filtered_df["iid"]==iid].iloc[0]
        new_title = st.text_input("Title", row["title"])
        new_desc = st.text_area("Description", row["description"])
        new_team = st.selectbox("Team", [""] + list(df["team"].unique()), index=0)
        new_status = st.selectbox("Status", [""] + list(df["status"].unique()), index=0)
        new_sprint = st.selectbox("Sprint", [""] + list(df["sprint"].unique()), index=0)
        new_project = st.selectbox("Project", [""] + list(df["project"].unique()), index=0)
        new_workstream = st.selectbox("Workstream", [""] + list(df["workstream"].unique()), index=0)
        new_milestone = st.selectbox("Milestone", [""] + list(df["milestone"].unique()), index=0)
        if st.button("Apply Changes"):
            labels=[]
            if new_team: labels.append(f"Team::{new_team}")
            if new_status: labels.append(f"Status::{new_status}")
            if new_sprint: labels.append(f"Sprint::{new_sprint}")
            if new_project: labels.append(f"Project::{new_project}")
            if new_workstream: labels.append(f"Workstream::{new_workstream}")
            if new_milestone: labels.append(f"Milestone::{new_milestone}")
            payload={"title":new_title,"description":new_desc,"labels":labels}
            if update_issue(base_url, project_id, token, iid, payload):
                st.success("Issue updated successfully!")
            else:
                st.error("Failed to update issue.")
